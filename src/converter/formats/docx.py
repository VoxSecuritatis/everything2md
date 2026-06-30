# ================================================================
# formats/docx
# ================================================================
# Objective:
#       Forward: convert .docx to .md using Mammoth for text/
#       headings/lists, python-docx for merged-cell table
#       inspection, and inline image extraction to the sidecar.
#       Reverse: parse the .md and reconstruct a best-effort .docx
#       using python-docx (headings, paragraphs, bold, italic,
#       tables, lists, images from sidecar).
# Inputs:
#       - forward: Path to .docx source file
#       - reverse: Path to .md file
# Outputs:
#       - forward: sibling .md + {stem}_assets/ sidecar
#       - reverse: sibling .docx
# Notes:
#       - Merged cells in tables are content-duplicated (Option A).
#       - Nested tables flattened to inline text.
#       - Image extraction failures emit warnings, do not fail the job.
# ================================================================

from __future__ import annotations

import re
import time
from pathlib import Path

from src.converter.assets import (
    asset_marker,
    ensure_images_dir,
    resolve_images_from_meta,
    write_meta,
)
from src.converter.naming import forward_output, reverse_output
from src.schemas import ConversionResult

# ------------------------------------------------
# Forward: DOCX -> MD
# ------------------------------------------------

def convert_forward(
    source_path: Path, options: dict, dest_dir: Path | None = None
) -> ConversionResult:
    """Convert a .docx file to a sibling .md file."""
    import mammoth  # type: ignore

    start = time.monotonic()
    output_path = forward_output(source_path, dest_dir)
    warnings: list[str] = []
    image_count = 0
    table_count = 0

    try:
        img_dir = ensure_images_dir(output_path)
        collector: list[str] = []

        image_converter = make_image_converter(img_dir, output_path.stem, collector)

        with source_path.open("rb") as fh:
            result = mammoth.convert_to_markdown(
                fh,
                convert_image=image_converter,
                style_map="p[style-name='List Paragraph'] => ol > li:fresh",
            )

        md_text = result.value
        image_count = len(collector)

        for msg in result.messages:
            if msg.type == "warning":
                warnings.append(msg.message)

        # Post-process: remove Word bookmark anchors, fix over-escaped chars
        md_text = postprocess_mammoth(md_text)

        # Table handling via python-docx for merged cells
        md_text, table_warnings, table_count = process_tables(source_path, md_text)
        warnings.extend(table_warnings)

        today = _today()
        front_matter = (
            f"---\n"
            f"source: {source_path.name}\n"
            f"date: {today}\n"
            f"tables: {table_count}\n"
            f"images: {image_count}\n"
            f"---\n"
        )
        marker = asset_marker(output_path)
        final_md = f"{front_matter}\n{marker}\n\n{md_text.strip()}\n"

        output_path.write_text(final_md, encoding="utf-8")
        write_meta(output_path, "docx", image_count=image_count, table_count=table_count)

        # Remove empty images dir if no images were extracted
        img_dir_path = img_dir
        if img_dir_path.exists() and not any(img_dir_path.iterdir()):
            img_dir_path.rmdir()
            img_dir_path.parent.rmdir() if not any(img_dir_path.parent.iterdir()) else None

    except Exception as exc:
        return ConversionResult(
            input_path=source_path,
            output_path=None,
            direction="forward",
            source_format="docx",
            status="failure",
            error=str(exc),
        )

    elapsed = (time.monotonic() - start) * 1000
    return ConversionResult(
        input_path=source_path,
        output_path=output_path,
        direction="forward",
        source_format="docx",
        status="success",
        warnings=warnings,
        image_count=image_count,
        table_count=table_count,
        duration_ms=elapsed,
    )


# ------------------------------------------------
# Reverse: MD -> DOCX
# ------------------------------------------------

def convert_reverse(
    md_path: Path, options: dict, dest_dir: Path | None = None
) -> ConversionResult:
    """Convert a .md file back to a best-effort .docx."""
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    start = time.monotonic()
    output_path = reverse_output(md_path, "docx", dest_dir)
    warnings: list[str] = []
    image_count = 0

    try:
        md_text = md_path.read_text(encoding="utf-8", errors="replace")
        img_dir = resolve_images_from_meta(md_path)
        lines = md_text.splitlines()
        doc = Document()

        i = 0
        while i < len(lines):
            line = lines[i]

            # Skip YAML front matter
            if i == 0 and line.strip() == "---":
                i += 1
                while i < len(lines) and lines[i].strip() != "---":
                    i += 1
                i += 1
                continue

            # Skip asset marker
            if line.strip().startswith("<!-- assets:"):
                i += 1
                continue

            # Fenced code block
            if line.startswith("```"):
                code_lines: list[str] = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                para = doc.add_paragraph("\n".join(code_lines))
                para.style = "Normal"
                run = para.runs[0] if para.runs else para.add_run("")
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                i += 1
                continue

            # Heading
            heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = strip_inline(heading_match.group(2))
                doc.add_heading(text, level=min(level, 9))
                i += 1
                continue

            # Image
            img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", line)
            if img_match and img_dir:
                rel_path = img_match.group(2)
                img_file = md_path.parent / rel_path
                if img_file.exists():
                    try:
                        doc.add_picture(str(img_file))
                        image_count += 1
                    except Exception as exc:
                        warnings.append(f"Could not embed image {img_file.name}: {exc}")
                i += 1
                continue

            # Unordered list item
            if re.match(r"^[-*+]\s+", line):
                text = strip_inline(re.sub(r"^[-*+]\s+", "", line))
                doc.add_paragraph(text, style="List Bullet")
                i += 1
                continue

            # Ordered list item
            if re.match(r"^\d+\.\s+", line):
                text = strip_inline(re.sub(r"^\d+\.\s+", "", line))
                doc.add_paragraph(text, style="List Number")
                i += 1
                continue

            # Table: collect table block, convert to docx table
            if line.startswith("|"):
                table_lines: list[str] = []
                while i < len(lines) and lines[i].startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                rows = parse_md_table(table_lines)
                if rows:
                    add_table_to_doc(doc, rows)
                continue

            # Blank line
            if not line.strip():
                i += 1
                continue

            # Regular paragraph (with inline formatting)
            para = doc.add_paragraph()
            add_inline_runs(para, line)
            i += 1

        doc.save(str(output_path))

    except Exception as exc:
        return ConversionResult(
            input_path=md_path,
            output_path=None,
            direction="reverse",
            source_format="docx",
            status="failure",
            error=str(exc),
        )

    elapsed = (time.monotonic() - start) * 1000
    return ConversionResult(
        input_path=md_path,
        output_path=output_path,
        direction="reverse",
        source_format="docx",
        status="success",
        warnings=warnings,
        image_count=image_count,
        duration_ms=elapsed,
    )


# ------------------------------------------------
# Image converter factory (for Mammoth)
# ------------------------------------------------

def make_image_converter(img_dir: Path, docname: str, collector: list[str]):
    """Return a Mammoth image converter that saves images to img_dir."""
    import mammoth  # type: ignore

    MIME_TO_EXT: dict[str, str] = {
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/gif": "gif",
        "image/bmp": "bmp",
        "image/tiff": "tiff",
        "image/webp": "webp",
        "image/svg+xml": "svg",
    }

    counter = [0]

    def convert_image(image) -> dict:
        counter[0] += 1
        ext = MIME_TO_EXT.get(image.content_type, "png")
        filename = f"{docname}_image_{counter[0]}.{ext}"
        dest = img_dir / filename

        try:
            with image.open() as img_bytes:
                dest.write_bytes(img_bytes.read())
            collector.append(filename)
            return {"src": f"images/{filename}"}
        except Exception as exc:
            return {"src": "", "alt": f"[image extraction failed: {exc}]"}

    return mammoth.images.img_element(convert_image)


# ------------------------------------------------
# Mammoth post-processing
# ------------------------------------------------

def postprocess_mammoth(md_text: str) -> str:
    """Remove Word artifacts from Mammoth output."""
    # Word internal bookmark anchors
    md_text = re.sub(r'<a id="_Hlk[^"]*"></a>', "", md_text)
    # Over-escaped punctuation: \. \( \) \-
    md_text = re.sub(r"\\([.()\-])", r"\1", md_text)
    # Collapse image alt text to single line (newlines break MD syntax)
    md_text = re.sub(r"!\[([^\]]*)\n([^\]]*)\]", lambda m: f"![{m.group(1)} {m.group(2)}]", md_text)
    # Remove bold markers wrapping images
    md_text = re.sub(r"\*\*(!\[[^\]]*\]\([^)]*\))\*\*", r"\1", md_text)
    # Drop empty bullet items
    md_text = re.sub(r"^[-*+]\s*$", "", md_text, flags=re.MULTILINE)
    # Collapse excess blank lines
    md_text = re.sub(r"\n{3,}", "\n\n", md_text)
    return md_text


# ------------------------------------------------
# Table handling (python-docx for merged cells)
# ------------------------------------------------

def process_tables(docx_path: Path, md_text: str) -> tuple[str, list[str], int]:
    """
    Re-render all tables from the .docx using python-docx (which handles
    merged cells correctly), then splice them back into md_text replacing
    Mammoth's table output. Returns (updated_md, warnings, table_count).
    """
    from docx import Document  # type: ignore

    warnings: list[str] = []
    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        return md_text, [f"Table processing skipped: {exc}"], 0

    if not doc.tables:
        return md_text, [], 0

    rendered_tables: list[str] = []
    for tbl_idx, table in enumerate(doc.tables):
        try:
            md_table, tbl_warnings = table_to_markdown(table, tbl_idx)
            rendered_tables.append(md_table)
            warnings.extend(tbl_warnings)
        except Exception as exc:
            warnings.append(f"Table {tbl_idx + 1}: could not render ({exc})")
            rendered_tables.append("")

    # Replace Mammoth-generated table blocks in document order
    md_text = replace_table_blocks(md_text, rendered_tables)
    return md_text, warnings, len(rendered_tables)


def build_cell_grid(table) -> list[list[str]]:
    """Build a 2D grid with merged cells duplicated (Option A)."""
    from docx.oxml.ns import qn  # type: ignore

    rows = table.rows
    if not rows:
        return []

    col_count = max(len(r.cells) for r in rows)
    grid: list[list[str]] = [[""] * col_count for _ in range(len(rows))]

    for row_idx, row in enumerate(rows):
        seen = []
        cells = row.cells
        col_idx = 0
        for cell in cells:
            # Horizontal merge detection: same _tc object appears multiple times
            if cell._tc in seen:
                grid[row_idx][col_idx] = grid[row_idx][col_idx - 1]
                col_idx += 1
                continue
            seen.append(cell._tc)

            # Vertical merge detection
            vmerge = cell._tc.find(qn("w:vMerge"))
            if vmerge is not None and vmerge.get(qn("w:val")) is None:
                # Continuation cell: copy from row above
                if row_idx > 0:
                    grid[row_idx][col_idx] = grid[row_idx - 1][col_idx]
                col_idx += 1
                continue

            text = cell_text(cell)
            grid[row_idx][col_idx] = text
            col_idx += 1
            if col_idx >= col_count:
                break

    return grid


def cell_text(cell) -> str:
    """Extract plain text from a cell, flattening nested tables."""
    parts: list[str] = []
    for block in cell.tables:
        nested_parts = []
        for row in block.rows:
            nested_parts.append(" / ".join(cell_text(c) for c in row.cells))
        parts.append(" | ".join(nested_parts))

    for para in cell.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    return " ".join(parts)


def table_to_markdown(table, tbl_idx: int) -> tuple[str, list[str]]:
    """Convert a python-docx Table to a GFM Markdown table string."""
    warnings: list[str] = []
    grid = build_cell_grid(table)
    if not grid:
        return "", []

    col_count = max(len(row) for row in grid)
    # Normalize row lengths
    for row in grid:
        while len(row) < col_count:
            row.append("")

    def escape(cell: str) -> str:
        return cell.replace("|", r"\|").replace("\n", " ")

    rows_md: list[str] = []
    for row_idx, row in enumerate(grid):
        cells = [escape(c) for c in row]
        rows_md.append("| " + " | ".join(cells) + " |")
        if row_idx == 0:
            rows_md.append("| " + " | ".join(["---"] * col_count) + " |")

    return "\n".join(rows_md), warnings


def replace_table_blocks(md_text: str, rendered_tables: list[str]) -> str:
    """Replace Mammoth table blocks with our re-rendered versions."""
    # Match pipe-table blocks (lines starting with |)
    table_block_pattern = re.compile(r"((?:^\|.+\n?)+)", re.MULTILINE)
    matches = list(table_block_pattern.finditer(md_text))

    # Replace from last to first to preserve offsets
    for idx, match in reversed(list(enumerate(matches))):
        if idx < len(rendered_tables) and rendered_tables[idx]:
            replacement = rendered_tables[idx] + "\n\n"
            md_text = md_text[: match.start()] + replacement + md_text[match.end():]

    # Append tables that Mammoth missed (idx >= len(matches))
    for idx in range(len(matches), len(rendered_tables)):
        if rendered_tables[idx]:
            md_text += "\n\n" + rendered_tables[idx] + "\n"

    return md_text


# ------------------------------------------------
# Reverse helpers
# ------------------------------------------------

def strip_inline(text: str) -> str:
    """Remove inline markdown from a single line."""
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.+?)_{1,3}", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    return text


def add_inline_runs(para, text: str) -> None:
    """Parse inline bold/italic/code in text and add runs to a paragraph."""
    token_pattern = re.compile(
        r"(\*{1,3}|\_{1,3}|`)"
        r"(.+?)"
        r"(\*{1,3}|\_{1,3}|`)"
    )
    last = 0
    for m in token_pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        marker = m.group(1)
        content = m.group(2)
        run = para.add_run(content)
        if marker in ("**", "__"):
            run.bold = True
        elif marker in ("***", "___"):
            run.bold = True
            run.italic = True
        elif marker in ("*", "_"):
            run.italic = True
        elif marker == "`":
            from docx.shared import Pt  # type: ignore
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parse a GFM table block into a list of row lists."""
    rows: list[list[str]] = []
    for line in lines:
        if re.match(r"^\|[-| :]+\|$", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows: list[list[str]]) -> None:
    """Add a python-docx Table to the document from parsed rows."""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < col_count:
                tbl.cell(r_idx, c_idx).text = strip_inline(cell_text)


def _today() -> str:
    from datetime import date
    return date.today().isoformat()
